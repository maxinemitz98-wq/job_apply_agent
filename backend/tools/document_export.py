"""
document_export — custom tool for agents to export markdown content to DOCX files.
The tool_schema dict is registered with the Anthropic API when building agent tool lists.
The execute() function is called by the backend when Claude emits a tool_use block.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


# Tool schema registered with Anthropic API
TOOL_SCHEMA = {
    "name": "document_export",
    "description": (
        "Exports markdown content to a formatted DOCX file. "
        "Use after writing the tailored CV or motivation letter. "
        "Returns the file path of the exported file."
    ),
    "input_schema": {
        "type": "object",
        "properties": {
            "content": {
                "type": "string",
                "description": "Markdown content to export",
            },
            "doc_type": {
                "type": "string",
                "enum": ["cv", "letter", "research_brief"],
                "description": "Document type — affects formatting",
            },
            "filename": {
                "type": "string",
                "description": "Output filename without extension",
            },
        },
        "required": ["content", "doc_type", "filename"],
    },
}


def execute(content: str, doc_type: str, filename: str, session_dir: Path) -> str:
    """
    Convert markdown to DOCX and save to session_dir.
    Returns the absolute file path of the saved DOCX.
    """
    session_dir.mkdir(parents=True, exist_ok=True)
    out_path = session_dir / f"{filename}.docx"

    doc = Document()
    _set_margins(doc)
    _apply_markdown(doc, content, doc_type)
    doc.save(str(out_path))

    return str(out_path)


def _set_margins(doc: Document):
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)


def _apply_markdown(doc: Document, content: str, doc_type: str):
    """Minimal markdown-to-docx renderer covering headings, bold, bullets, paragraphs."""
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]

        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline(p, line[2:])
        elif line.strip() == "":
            pass  # blank line — skip
        else:
            p = doc.add_paragraph()
            _add_inline(p, line)

        i += 1


def _add_inline(paragraph, text: str):
    """Handle **bold** inline formatting."""
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            paragraph.add_run(part)
